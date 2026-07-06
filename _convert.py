from docx import Document
import re

doc = Document()

with open('papers/混沌的结构学解释.md', 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

for line in lines:
    line = line.strip()
    if not line: continue
    if line.startswith('# ') and '##' not in line:
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'): continue
    elif line.startswith('> '):
        doc.add_paragraph(line[2:])
    elif re.match(r'^\|.*\|$', line):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= set('-: ') for c in cells): continue
        doc.add_paragraph(' | '.join(cells))
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif re.match(r'^\d+\.', line):
        doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        p.add_run(line[2:-2]).bold = True
    else:
        doc.add_paragraph(line)

doc.save('papers/混沌的结构学解释.docx')
print('done')
